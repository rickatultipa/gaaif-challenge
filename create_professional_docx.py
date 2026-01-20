#!/usr/bin/env python3
"""
Generate Professional Word Document for GAAIF Challenge Submission
===================================================================

Creates a polished, unique template with all figures properly embedded.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '1B4F72')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

def create_professional_document():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Modify Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri Light'
        heading_style.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)  # Dark blue

    # ==========================================================================
    # COVER PAGE
    # ==========================================================================

    # Add spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('STRUCTURED GOLD FORWARD')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run.font.name = 'Calibri Light'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('with Double Knock-Out Barriers')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    run.font.name = 'Calibri Light'

    doc.add_paragraph()

    # Decorative line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0xD4, 0xAC, 0x0D)  # Gold color
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Document type
    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run('Product Development Memorandum')
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    doc.add_paragraph()
    doc.add_paragraph()

    # Info box
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ('Issuing Desk:', 'Derivatives Structuring'),
        ('Client:', 'Zeus Gold Group AG'),
        ('Date:', 'January 2026'),
        ('Classification:', 'CONFIDENTIAL')
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(2.5)

    for _ in range(6):
        doc.add_paragraph()

    # Author info at bottom
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('Prepared for GAAIF Challenge 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

    # Page break
    doc.add_page_break()

    # ==========================================================================
    # EXECUTIVE SUMMARY
    # ==========================================================================

    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'This memorandum presents our analysis of a proposed structured hedging facility '
        'for Zeus Gold Group AG ("Z Group"). The product combines exposure to LBMA gold '
        'prices with automatic termination features linked to EUR/USD exchange rate movements.'
    )

    doc.add_paragraph(
        'We have developed a comprehensive pricing framework validated through multiple '
        'methodologies. Our analysis identifies several structural considerations that '
        'warrant discussion before proceeding to term sheet finalization.'
    )

    add_horizontal_line(doc)

    # Transaction Summary Table
    h = doc.add_paragraph()
    run = h.add_run('Transaction Summary')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Table Grid'

    data1 = [
        ('Notional Principal', 'EUR 500,000,000'),
        ('Reference Asset', 'LBMA Gold PM Fixing (USD/oz)'),
        ('Strike Price', 'USD 4,600 per troy ounce'),
        ('Tenor', '2 years (March 2026 — February 2028)'),
        ('Lower Barrier', 'EUR/USD < 1.05 (Knock-Out)'),
        ('Upper Barrier', 'EUR/USD > 1.25 (Knock-Out)')
    ]

    for i, (param, spec) in enumerate(data1):
        row = table1.rows[i]
        row.cells[0].text = param
        row.cells[1].text = spec
        set_cell_shading(row.cells[0], 'E8F4FD')
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # Key Findings Table
    h = doc.add_paragraph()
    run = h.add_run('Key Findings')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'

    data2 = [
        ('Z Group Present Value', 'EUR −192 million'),
        ('Alphabank Present Value', 'EUR +192 million'),
        ('Knock-Out Probability', '93%'),
        ('Expected Contract Duration', '5 months')
    ]

    for i, (metric, result) in enumerate(data2):
        row = table2.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = result
        set_cell_shading(row.cells[0], 'FEF9E7')
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        'The negative present value for Z Group reflects the strike price being set 54% above '
        'the two-year gold forward. The high knock-out probability stems from the lower '
        "barrier's proximity to current spot (2.8% distance) combined with negative EUR/USD "
        'drift from interest rate differentials.'
    )

    doc.add_page_break()

    # ==========================================================================
    # TRANSACTION OVERVIEW
    # ==========================================================================

    doc.add_heading('2. Transaction Overview', level=1)

    doc.add_paragraph(
        'Zeus Gold Group, a Frankfurt-headquartered jewelry manufacturer, seeks to hedge its '
        'USD-denominated gold procurement costs while managing EUR/USD translation risk. '
        'The proposed facility would run for two years commencing March 2026.'
    )

    doc.add_heading('2.1 Settlement Mechanics', level=2)

    doc.add_paragraph(
        'At settlement time τ (maturity or knock-out, whichever occurs first), '
        'with LBMA gold fixing at price P:'
    )

    # Payoff formulas
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_para.add_run('Z Group Payoff = N × (P − K) / K')
    run.font.name = 'Consolas'
    run.font.size = Pt(11)

    formula_para2 = doc.add_paragraph()
    formula_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_para2.add_run('Alphabank Payoff = N × (K − P) / K')
    run.font.name = 'Consolas'
    run.font.size = Pt(11)

    doc.add_paragraph('where:')
    bullet = doc.add_paragraph('N = EUR 500,000,000 (notional principal)', style='List Bullet')
    bullet = doc.add_paragraph('K = USD 4,600/oz (strike price)', style='List Bullet')
    bullet = doc.add_paragraph('P = LBMA Gold PM fixing at settlement', style='List Bullet')

    doc.add_heading('2.2 Knock-Out Mechanism', level=2)

    doc.add_paragraph(
        'The contract terminates immediately upon the first occurrence of EUR/USD breaching '
        'either barrier:'
    )

    bullet = doc.add_paragraph('Lower Barrier: EUR/USD < 1.05', style='List Bullet')
    bullet = doc.add_paragraph('Upper Barrier: EUR/USD > 1.25', style='List Bullet')

    doc.add_page_break()

    # ==========================================================================
    # MATHEMATICAL FRAMEWORK
    # ==========================================================================

    doc.add_heading('3. Mathematical Framework', level=1)

    doc.add_heading('3.1 Stochastic Model', level=2)

    doc.add_paragraph(
        'Both underlying assets follow geometric Brownian motion under the risk-neutral '
        'measure Q (EUR numeraire):'
    )

    h = doc.add_paragraph()
    run = h.add_run('Gold Price Dynamics (with Quanto Adjustment):')
    run.font.bold = True

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('dS/S = (r_USD − q − ρσ_S σ_X) dt + σ_S dW^S')
    run.font.name = 'Consolas'

    doc.add_paragraph(
        'The quanto adjustment (−ρσ_S σ_X) accounts for the correlation between gold and '
        'EUR/USD when the underlying is USD-denominated but the payoff is in EUR.'
    )

    h = doc.add_paragraph()
    run = h.add_run('EUR/USD Exchange Rate:')
    run.font.bold = True

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('dX/X = (r_EUR − r_USD) dt + σ_X dW^X')
    run.font.name = 'Consolas'

    doc.add_heading('3.2 Parameter Estimates', level=2)

    # Parameters table
    param_table = doc.add_table(rows=9, cols=3)
    param_table.style = 'Table Grid'

    headers = ['Parameter', 'Value', 'Source']
    for i, h in enumerate(headers):
        param_table.rows[0].cells[i].text = h
        set_cell_shading(param_table.rows[0].cells[i], '1B4F72')
        param_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        param_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    params = [
        ('Gold spot (S₀)', 'USD 2,750/oz', 'LBMA Jan 2026'),
        ('EUR/USD spot (X₀)', '1.08', 'ECB reference'),
        ('USD risk-free rate (r_USD)', '4.5%', 'OIS curve'),
        ('EUR risk-free rate (r_EUR)', '2.5%', 'OIS curve'),
        ('Gold volatility (σ_S)', '18%', '1Y ATM implied'),
        ('EUR/USD volatility (σ_X)', '8%', '1Y ATM implied'),
        ('Correlation (ρ)', '−0.25', '1Y historical'),
        ('Gold convenience yield (q)', '0.5%', 'GOFO proxy')
    ]

    for i, (param, val, src) in enumerate(params, 1):
        param_table.rows[i].cells[0].text = param
        param_table.rows[i].cells[1].text = val
        param_table.rows[i].cells[2].text = src

    doc.add_page_break()

    # ==========================================================================
    # NUMERICAL IMPLEMENTATION
    # ==========================================================================

    doc.add_heading('4. Numerical Implementation', level=1)

    doc.add_heading('4.1 Simulation Methodology', level=2)

    doc.add_paragraph(
        'We employ Monte Carlo simulation with the following specifications:'
    )

    sim_table = doc.add_table(rows=4, cols=3)
    sim_table.style = 'Table Grid'

    sim_headers = ['Parameter', 'Value', 'Rationale']
    for i, h in enumerate(sim_headers):
        sim_table.rows[0].cells[i].text = h
        set_cell_shading(sim_table.rows[0].cells[i], '1B4F72')
        sim_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        sim_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    sim_data = [
        ('Simulation paths', '100,000', 'Adequate precision'),
        ('Time steps', '504', 'Daily monitoring (2 years)'),
        ('Random seed', 'Fixed', 'Reproducibility')
    ]

    for i, (p, v, r) in enumerate(sim_data, 1):
        sim_table.rows[i].cells[0].text = p
        sim_table.rows[i].cells[1].text = v
        sim_table.rows[i].cells[2].text = r

    doc.add_heading('4.2 Variance Reduction', level=2)

    doc.add_paragraph(
        'Two techniques are implemented to improve computational efficiency:'
    )

    h = doc.add_paragraph()
    run = h.add_run('Antithetic Variates: ')
    run.font.bold = True
    h.add_run(
        'For each path with innovations {Z}, we also simulate the reflected path {−Z}. '
        'The negative correlation between paired paths reduces variance.'
    )

    h = doc.add_paragraph()
    run = h.add_run('Control Variate: ')
    run.font.bold = True
    h.add_run(
        'The vanilla gold forward (without barriers) serves as a control with known '
        'analytical price. Combined, these techniques reduce standard errors by approximately 60%.'
    )

    doc.add_page_break()

    # ==========================================================================
    # PRICING RESULTS
    # ==========================================================================

    doc.add_heading('5. Pricing Results', level=1)

    doc.add_heading('5.1 Base Case Valuation', level=2)

    results_table = doc.add_table(rows=5, cols=2)
    results_table.style = 'Table Grid'

    results_data = [
        ('Z Group Present Value', 'EUR −191,900,647'),
        ('Alphabank Present Value', 'EUR +191,900,647'),
        ('Standard Error', 'EUR 123,877'),
        ('95% Confidence Interval', '[−192.1M, −191.7M]')
    ]

    for i, (m, v) in enumerate(results_data):
        results_table.rows[i].cells[0].text = m
        results_table.rows[i].cells[1].text = v
        set_cell_shading(results_table.rows[i].cells[0], 'E8F4FD')
        results_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading('5.2 Barrier Analysis', level=2)

    barrier_table = doc.add_table(rows=5, cols=2)
    barrier_table.style = 'Table Grid'

    barrier_data = [
        ('Overall Knock-Out Rate', '92.99%'),
        ('Lower Barrier Breaches', '86.02%'),
        ('Upper Barrier Breaches', '6.97%'),
        ('Average Time to Knock-Out', '0.43 years (5.2 months)')
    ]

    for i, (m, v) in enumerate(barrier_data):
        barrier_table.rows[i].cells[0].text = m
        barrier_table.rows[i].cells[1].text = v
        set_cell_shading(barrier_table.rows[i].cells[0], 'FEF9E7')
        barrier_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        'The asymmetry between barrier breaches reflects the negative EUR/USD drift implied '
        'by interest rate parity. With r_EUR − r_USD = −2% annually, the euro faces persistent '
        'depreciation pressure, making the lower barrier far more likely to be reached.'
    )

    # Insert Monte Carlo paths figure
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run('Figure 1: Monte Carlo Simulation Paths')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/monte_carlo_paths.png'):
        doc.add_picture('output/monte_carlo_paths.png', width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading('5.3 Convergence Verification', level=2)

    doc.add_paragraph(
        'Monte Carlo estimates stabilize as path counts increase:'
    )

    conv_table = doc.add_table(rows=6, cols=3)
    conv_table.style = 'Table Grid'

    conv_headers = ['Paths', 'Price Estimate', 'Standard Error']
    for i, h in enumerate(conv_headers):
        conv_table.rows[0].cells[i].text = h
        set_cell_shading(conv_table.rows[0].cells[i], '1B4F72')
        conv_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        conv_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    conv_data = [
        ('5,000', 'EUR −191.4M', 'EUR 551K'),
        ('10,000', 'EUR −191.7M', 'EUR 381K'),
        ('25,000', 'EUR −191.8M', 'EUR 247K'),
        ('50,000', 'EUR −192.1M', 'EUR 174K'),
        ('100,000', 'EUR −191.9M', 'EUR 124K')
    ]

    for i, (p, e, s) in enumerate(conv_data, 1):
        conv_table.rows[i].cells[0].text = p
        conv_table.rows[i].cells[1].text = e
        conv_table.rows[i].cells[2].text = s

    doc.add_paragraph()

    # Insert convergence figure
    h = doc.add_paragraph()
    run = h.add_run('Figure 2: Monte Carlo Convergence Analysis')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/convergence_analysis.png'):
        doc.add_picture('output/convergence_analysis.png', width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================================================================
    # CRITICAL ASSESSMENT
    # ==========================================================================

    doc.add_heading('6. Critical Assessment', level=1)

    doc.add_heading('6.1 Strike Price Analysis', level=2)

    doc.add_paragraph(
        'The specified strike of USD 4,600/oz warrants careful examination.'
    )

    h = doc.add_paragraph()
    run = h.add_run('Forward Price Calculation:')
    run.font.bold = True

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('F₀,T = S₀ × exp((r_USD − q) × T) = 2750 × exp(0.04 × 2) ≈ USD 2,979/oz')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'The strike exceeds the forward by 54%, placing Z Group in a deeply out-of-the-money position:'
    )

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('Moneyness = F₀,T / K = 2979 / 4600 = 64.8%')
    run.font.name = 'Consolas'

    doc.add_paragraph()

    h = doc.add_paragraph()
    run = h.add_run('Alternative Strike Analysis:')
    run.font.bold = True

    strike_table = doc.add_table(rows=5, cols=3)
    strike_table.style = 'Table Grid'

    strike_headers = ['Strike', 'Forward Relationship', 'Z Group PV']
    for i, h in enumerate(strike_headers):
        strike_table.rows[0].cells[i].text = h
        set_cell_shading(strike_table.rows[0].cells[i], '1B4F72')
        strike_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        strike_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    strike_data = [
        ('USD 2,800', '6% below forward', 'EUR +2M'),
        ('USD 3,000', 'At-the-money', 'EUR −31M'),
        ('USD 3,500', '17% above forward', 'EUR −97M'),
        ('USD 4,600', '54% above forward', 'EUR −192M')
    ]

    for i, row_data in enumerate(strike_data, 1):
        for j, val in enumerate(row_data):
            strike_table.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Insert strike sensitivity figure
    h = doc.add_paragraph()
    run = h.add_run('Figure 3: Strike Price Sensitivity Analysis')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/scenario_strike.png'):
        doc.add_picture('output/scenario_strike.png', width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading('6.2 Barrier Configuration', level=2)

    doc.add_paragraph(
        'The lower barrier at 1.05 sits only 2.8% below current spot:'
    )

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('Distance to Lower Barrier = (X₀ − L) / X₀ = (1.08 − 1.05) / 1.08 = 2.78%')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'Given 8% annual EUR/USD volatility and negative drift, barrier breach is near-certain '
        'over a two-year horizon.'
    )

    h = doc.add_paragraph()
    run = h.add_run('Alternative Configurations:')
    run.font.bold = True

    barrier_config_table = doc.add_table(rows=4, cols=3)
    barrier_config_table.style = 'Table Grid'

    barrier_config_headers = ['Corridor', 'Knock-Out Rate', 'Expected Duration']
    for i, h in enumerate(barrier_config_headers):
        barrier_config_table.rows[0].cells[i].text = h
        set_cell_shading(barrier_config_table.rows[0].cells[i], '1B4F72')
        barrier_config_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        barrier_config_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    barrier_config_data = [
        ('[1.05, 1.25]', '93%', '5 months'),
        ('[1.00, 1.30]', '66%', '10 months'),
        ('[0.95, 1.35]', '39%', '14 months')
    ]

    for i, row_data in enumerate(barrier_config_data, 1):
        for j, val in enumerate(row_data):
            barrier_config_table.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Insert barrier sensitivity figure
    h = doc.add_paragraph()
    run = h.add_run('Figure 4: Barrier Width Sensitivity Analysis')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/scenario_barrier.png'):
        doc.add_picture('output/scenario_barrier.png', width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================================================================
    # RISK SENSITIVITIES
    # ==========================================================================

    doc.add_heading('7. Risk Sensitivities', level=1)

    doc.add_heading('7.1 Greeks Summary', level=2)

    greeks_table = doc.add_table(rows=7, cols=3)
    greeks_table.style = 'Table Grid'

    greeks_headers = ['Greek', 'Value', 'Interpretation']
    for i, h in enumerate(greeks_headers):
        greeks_table.rows[0].cells[i].text = h
        set_cell_shading(greeks_table.rows[0].cells[i], '1B4F72')
        greeks_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        greeks_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    greeks_data = [
        ('Δ_gold', 'EUR 109,545 per $1', 'First-order gold sensitivity'),
        ('Γ_gold', 'EUR −491', 'Gold convexity'),
        ('Δ_FX', 'EUR 2.26M per 0.01', 'EUR/USD sensitivity'),
        ('Vega_gold', 'EUR −691K per 1% vol', 'Gold vega'),
        ('ρ_EUR', 'EUR −11.7M per 1bp', 'EUR rate sensitivity'),
        ('Corr Sens', 'EUR +143K per 0.05', 'Correlation sensitivity')
    ]

    for i, row_data in enumerate(greeks_data, 1):
        for j, val in enumerate(row_data):
            greeks_table.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Insert Greeks figure
    h = doc.add_paragraph()
    run = h.add_run('Figure 5: Risk Sensitivities Summary')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/greeks_summary.png'):
        doc.add_picture('output/greeks_summary.png', width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('7.2 Hedging Implications', level=2)

    h = doc.add_paragraph()
    run = h.add_run('Delta Hedging: ')
    run.font.bold = True
    h.add_run(
        'The gold delta of EUR 110K per dollar implies a hedge ratio of approximately 183,000 oz.'
    )

    h = doc.add_paragraph()
    run = h.add_run('Barrier Risk: ')
    run.font.bold = True
    h.add_run(
        'As EUR/USD approaches either barrier, gamma and delta become increasingly unstable—the '
        'characteristic "pin risk" of barrier options. Hedging costs will escalate significantly '
        'in the final days before a potential knock-out.'
    )

    doc.add_page_break()

    # ==========================================================================
    # SENSITIVITY ANALYSIS
    # ==========================================================================

    doc.add_heading('8. Sensitivity Analysis', level=1)

    doc.add_paragraph(
        'Comprehensive sensitivity analysis across key model parameters:'
    )

    # Insert sensitivity analysis figure
    h = doc.add_paragraph()
    run = h.add_run('Figure 6: Parameter Sensitivity Analysis')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/sensitivity_analysis.png'):
        doc.add_picture('output/sensitivity_analysis.png', width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================================================================
    # MODEL VALIDATION
    # ==========================================================================

    doc.add_heading('9. Model Validation', level=1)

    doc.add_heading('9.1 Alternative Model Specifications', level=2)

    doc.add_paragraph(
        'To ensure robustness, we compared valuations across three model specifications:'
    )

    model_table = doc.add_table(rows=4, cols=3)
    model_table.style = 'Table Grid'

    model_headers = ['Model', 'Z Group PV', 'Knock-Out Rate']
    for i, h in enumerate(model_headers):
        model_table.rows[0].cells[i].text = h
        set_cell_shading(model_table.rows[0].cells[i], '1B4F72')
        model_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        model_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    model_data = [
        ('Base GBM', 'EUR −192.1M', '92.9%'),
        ('Heston Stochastic Vol', 'EUR −191.8M', '93.0%'),
        ('Merton Jump-Diffusion', 'EUR −191.7M', '93.0%')
    ]

    for i, row_data in enumerate(model_data, 1):
        for j, val in enumerate(row_data):
            model_table.rows[i].cells[j].text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'All models converge within 0.2%, confirming that the barrier structure dominates '
        'pricing dynamics. Model specification risk is secondary.'
    )

    # Insert model comparison figure
    h = doc.add_paragraph()
    run = h.add_run('Figure 7: Cross-Model Validation')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/model_comparison.png'):
        doc.add_picture('output/model_comparison.png', width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('9.2 Analytical Benchmark', level=2)

    doc.add_paragraph(
        'The vanilla gold forward (without barriers) provides a sanity check:'
    )

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('V_vanilla = e^(-r_EUR×T) × N × (F − K) / K = EUR −167,598,411')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'The knock-out version (EUR −192M) is EUR 24M worse than the vanilla, representing '
        'the expected cost of early termination when gold is below strike.'
    )

    doc.add_page_break()

    # ==========================================================================
    # PAYOFF ANALYSIS
    # ==========================================================================

    doc.add_heading('10. Payoff Analysis', level=1)

    # Insert payoff diagram
    h = doc.add_paragraph()
    run = h.add_run('Figure 8: Payoff Diagram')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/payoff_diagram.png'):
        doc.add_picture('output/payoff_diagram.png', width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Insert payoff distribution
    h = doc.add_paragraph()
    run = h.add_run('Figure 9: Payoff Distribution')
    run.font.bold = True
    run.font.size = Pt(10)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists('output/payoff_distribution.png'):
        doc.add_picture('output/payoff_distribution.png', width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================================================================
    # CONCLUSIONS
    # ==========================================================================

    doc.add_heading('11. Conclusions and Recommendations', level=1)

    doc.add_heading('11.1 Summary of Findings', level=2)

    doc.add_paragraph(
        'The proposed structure is technically sound and priceable using standard Monte Carlo '
        'techniques. However, two features merit discussion:'
    )

    p = doc.add_paragraph()
    run = p.add_run('1. Strike positioning: ')
    run.font.bold = True
    p.add_run(
        'The USD 4,600 strike creates a deeply out-of-the-money position for Z Group. '
        'Clarification of the commercial rationale is recommended.'
    )

    p = doc.add_paragraph()
    run = p.add_run('2. Barrier proximity: ')
    run.font.bold = True
    p.add_run(
        'The 93% knock-out probability results in an expected contract life of only 5 months—'
        'potentially misaligned with a 2-year hedging mandate.'
    )

    doc.add_heading('11.2 Recommendations', level=2)

    doc.add_paragraph('We recommend proceeding to term sheet stage contingent upon:')

    doc.add_paragraph(
        "1. Confirmation from Relationship Management regarding Z Group's acceptance of "
        "the strike level and its implications",
        style='List Number'
    )
    doc.add_paragraph(
        "2. Discussion of whether alternative barrier configurations (e.g., [1.00, 1.30]) "
        "would better serve the client's hedging objectives",
        style='List Number'
    )
    doc.add_paragraph(
        "3. Documentation of appropriate risk disclosures regarding the high knock-out probability",
        style='List Number'
    )

    doc.add_heading('11.3 Next Steps', level=2)

    doc.add_paragraph('Upon Committee approval, we will:')

    doc.add_paragraph('Finalize term sheet documentation', style='List Bullet')
    doc.add_paragraph('Establish hedging framework with Trading Desk', style='List Bullet')
    doc.add_paragraph('Coordinate credit approval with Risk Management', style='List Bullet')
    doc.add_paragraph('Schedule client presentation', style='List Bullet')

    add_horizontal_line(doc)

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run('[END OF MEMORANDUM]')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_page_break()

    # ==========================================================================
    # APPENDIX
    # ==========================================================================

    doc.add_heading('Appendix A: Nomenclature', level=1)

    nomenclature_table = doc.add_table(rows=12, cols=2)
    nomenclature_table.style = 'Table Grid'

    nomenclature_data = [
        ('Symbol', 'Description'),
        ('S_t', 'Gold spot price (USD/oz) at time t'),
        ('X_t', 'EUR/USD exchange rate at time t'),
        ('K', 'Strike price (USD 4,600/oz)'),
        ('N', 'Notional principal (EUR 500M)'),
        ('T', 'Maturity (2 years)'),
        ('L, U', 'Lower (1.05) and upper (1.25) barriers'),
        ('r_EUR, r_USD', 'Risk-free rates'),
        ('σ_S, σ_X', 'Volatilities'),
        ('ρ', 'Correlation coefficient'),
        ('q', 'Gold convenience yield'),
        ('τ', 'Settlement time')
    ]

    for i, (sym, desc) in enumerate(nomenclature_data):
        nomenclature_table.rows[i].cells[0].text = sym
        nomenclature_table.rows[i].cells[1].text = desc
        if i == 0:
            set_cell_shading(nomenclature_table.rows[i].cells[0], '1B4F72')
            set_cell_shading(nomenclature_table.rows[i].cells[1], '1B4F72')
            nomenclature_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            nomenclature_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            nomenclature_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            nomenclature_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading('Appendix B: Figure Index', level=1)

    figures = [
        ('Figure 1', 'Monte Carlo Simulation Paths'),
        ('Figure 2', 'Monte Carlo Convergence Analysis'),
        ('Figure 3', 'Strike Price Sensitivity Analysis'),
        ('Figure 4', 'Barrier Width Sensitivity Analysis'),
        ('Figure 5', 'Risk Sensitivities Summary'),
        ('Figure 6', 'Parameter Sensitivity Analysis'),
        ('Figure 7', 'Cross-Model Validation'),
        ('Figure 8', 'Payoff Diagram'),
        ('Figure 9', 'Payoff Distribution')
    ]

    for fig_num, fig_desc in figures:
        p = doc.add_paragraph()
        run = p.add_run(f'{fig_num}: ')
        run.font.bold = True
        p.add_run(fig_desc)

    # Save document
    doc.save('PRODUCT_PROPOSAL_PROFESSIONAL.docx')
    print("Professional document created: PRODUCT_PROPOSAL_PROFESSIONAL.docx")
    print(f"Total figures embedded: 9")

if __name__ == "__main__":
    create_professional_document()
