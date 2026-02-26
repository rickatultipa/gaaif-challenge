#!/usr/bin/env python3
"""
Generate PRODUCT_PROPOSAL_FINAL.docx from scratch using python-docx.
All values sourced from the pricing model with live/fallback market data.
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1A, 0x36, 0x5D)
HEADING_COLOR = RGBColor(0x0F, 0x47, 0x61)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = "F2F2F2"
MED_GRAY = "E0E0E0"


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg_hex="1A365D", font_color=WHITE):
    """Style a header row with navy background and white text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = font_color
                run.font.bold = True
                run.font.size = Pt(10)


def style_alternating_rows(table, start_row=1):
    """Apply alternating row colors to a table."""
    for i, row in enumerate(table.rows[start_row:], start=start_row):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    style_header_row(table.rows[0])

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

    style_alternating_rows(table)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table


def add_body_text(doc, text, bold=False, italic=False, alignment=None):
    """Add a body paragraph with consistent 11pt formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    return p


def add_formula(doc, text):
    """Add a formula line (indented, italic, 10.5pt)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Cambria Math'
    run.italic = True
    return p


def add_chart(doc, filename, width=Inches(6.0)):
    """Insert a chart image if it exists, centered."""
    path = os.path.join('output', 'pdf_charts', filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        add_body_text(doc, f"[Chart not found: {filename}]", italic=True)


def add_heading_numbered(doc, number, title, level=1):
    """Add a numbered heading in the document style."""
    h = doc.add_heading(level=level)
    run = h.add_run(f"{number}\t{title}")
    run.font.color.rgb = HEADING_COLOR
    return h


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # -- Default font --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading styles
    for level in [1, 2, 3]:
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = HEADING_COLOR
        hs.font.name = 'Calibri'
        if level == 1:
            hs.font.size = Pt(20)
            # Remove default "page break before" so content flows naturally
            hs.paragraph_format.page_break_before = False
        elif level == 2:
            hs.font.size = Pt(16)
        else:
            hs.font.size = Pt(13)

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Structured Gold Forward\nwith Knock-Out Barriers")
    run.font.size = Pt(28)
    run.font.color.rgb = HEADING_COLOR
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Product Development Memorandum")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Derivatives Structuring Desk")
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("February 2026")
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared By: Ricky Sun / ULTIPA")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # =========================================================================
    # MEMO HEADER
    # =========================================================================
    p = doc.add_paragraph()
    run = p.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    memo_fields = [
        ("TO:", "Product Committee, Alphabank S.A."),
        ("FROM:", "Derivatives Structuring Desk"),
        ("DATE:", "February 2026"),
        ("RE:", "Pricing and Risk Analysis \u2014 Zeus Gold Group AG Hedging Facility"),
    ]
    for label, value in memo_fields:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)

    # Horizontal rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/></w:pBdr>')
    pPr.append(pBdr)

    # Market Update callout (matching old doc's orange #C05000 style)
    p = doc.add_paragraph()
    run = p.add_run("MARKET UPDATE \u2014 February 26, 2026")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

    add_body_text(doc,
        "This memorandum has been updated with market data as of February 26, 2026. "
        "Gold has moved to $5,203/oz (from $4,900 on Feb 1), EUR/USD to 1.181, and "
        "realized gold volatility has risen to 41%. The position reversal is now firmly "
        "established with Z Group PV at EUR +64M.")

    doc.add_paragraph()

    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading("Executive Summary", level=1)

    add_body_text(doc,
        "This memorandum presents our analysis of a proposed structured hedging facility "
        "for Zeus Gold Group AG (\u201cZ Group\u201d). The product combines exposure to LBMA gold "
        "prices with automatic termination features linked to EUR/USD exchange rate movements.")

    add_body_text(doc,
        "We have developed a comprehensive pricing framework validated through multiple "
        "methodologies. Our analysis identifies several structural considerations that "
        "warrant discussion before proceeding to term sheet finalization.")

    # Transaction Summary table
    add_body_text(doc, "Transaction Summary", bold=True)
    add_styled_table(doc,
        ["Parameter", "Specification"],
        [
            ["Notional Principal", "EUR 500,000,000"],
            ["Reference Asset", "LBMA Gold PM Fixing (USD/oz)"],
            ["Strike Price", "USD 4,600 per troy ounce"],
            ["Tenor", "2 years (March 2026 \u2014 February 2028)"],
            ["Knock-Out Barriers", "EUR/USD < 1.05 or EUR/USD > 1.25"],
            ["Settlement", "Cash settled in EUR at prevailing EUR/USD rate"],
        ],
        col_widths=[2.5, 4.0]
    )

    doc.add_paragraph()

    # Key Findings table
    add_body_text(doc, "Key Findings", bold=True)
    add_styled_table(doc,
        ["Metric", "Result"],
        [
            ["Z Group Present Value", "EUR +64 million"],
            ["Alphabank Present Value", "EUR \u221264 million"],
            ["Knock-Out Probability", "64%"],
            ["Expected Contract Duration", "12 months"],
            ["Forward Moneyness", "113.1%"],
        ],
        col_widths=[3.0, 3.5]
    )

    doc.add_paragraph()
    add_body_text(doc,
        "The positive present value for Z Group reflects gold\u2019s surge above the strike price, "
        "placing the forward at 113% moneyness. The 64% knock-out probability is driven primarily "
        "by the upper barrier\u2019s proximity to current EUR/USD spot (5.8% distance), representing "
        "a fundamental reversal from earlier market conditions where the lower barrier dominated.")

    # =========================================================================
    # SECTION 1: TRANSACTION OVERVIEW
    # =========================================================================
    add_heading_numbered(doc, "1", "Transaction Overview")

    add_body_text(doc,
        "Zeus Gold Group, a Frankfurt-headquartered jewelry manufacturer, seeks to hedge its "
        "USD-denominated gold procurement costs while managing EUR/USD translation risk. The "
        "proposed facility would run for two years commencing March 2026.")

    # 1.1 Settlement Mechanics
    add_heading_numbered(doc, "1.1", "Settlement Mechanics", level=2)

    add_body_text(doc,
        "At settlement time \u03c4 (maturity or knock-out, whichever occurs first), with LBMA "
        "gold fixing at price P:")

    add_formula(doc, "Z Group Payoff = N \u00d7 (P\u03c4 \u2212 K) / K")
    add_formula(doc, "Alphabank Payoff = N \u00d7 (K \u2212 P\u03c4) / K")

    add_body_text(doc, "where:")
    add_body_text(doc, "\u2022  N = 500,000,000 EUR (notional principal)")
    add_body_text(doc, "\u2022  K = 4,600 USD/oz (strike price)")
    add_body_text(doc, "\u2022  P\u03c4 = LBMA Gold PM fixing at settlement")

    # 1.2 Knock-Out Mechanism
    add_heading_numbered(doc, "1.2", "Knock-Out Mechanism", level=2)

    add_body_text(doc,
        "The contract terminates immediately upon the first occurrence of:")

    add_formula(doc, "X\u209c < L  or  X\u209c > U")

    add_body_text(doc,
        "where X\u209c denotes the EUR/USD rate, L = 1.05 (lower barrier), and U = 1.25 (upper barrier).")

    add_body_text(doc, "The stopping time is defined as:")
    add_formula(doc, "\u03c4\u2096\u2092 = inf{t \u2265 0 : X\u209c \u2209 (L, U)}")

    add_body_text(doc,
        "Settlement occurs at \u03c4 = min(T, \u03c4\u2096\u2092) where T = 2 years.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 2: MATHEMATICAL FRAMEWORK
    # =========================================================================
    add_heading_numbered(doc, "2", "Mathematical Framework")

    # 2.1 Stochastic Model
    add_heading_numbered(doc, "2.1", "Stochastic Model", level=2)

    add_body_text(doc,
        "Both underlying assets follow geometric Brownian motion under the risk-neutral "
        "measure \u211a.")

    add_body_text(doc, "Gold Price Dynamics (USD):", bold=True)
    add_formula(doc, "dS\u209c/S\u209c = (r_USD \u2212 q) dt + \u03c3_S dW\u209c\u207d\u00b9\u207e")

    add_body_text(doc, "EUR/USD Exchange Rate:", bold=True)
    add_formula(doc, "dX\u209c/X\u209c = (r_EUR \u2212 r_USD) dt + \u03c3_X dW\u209c\u207d\u00b2\u207e")

    add_body_text(doc, "Correlation Structure:", bold=True)
    add_formula(doc, "\u1d3c[dW\u209c\u207d\u00b9\u207e \u00b7 dW\u209c\u207d\u00b2\u207e] = \u03c1 dt")

    # 2.2 Parameter Estimates
    add_heading_numbered(doc, "2.2", "Parameter Estimates", level=2)

    add_styled_table(doc,
        ["Parameter", "Symbol", "Value", "Source"],
        [
            ["Gold spot", "S\u2080", "USD 5,203/oz", "yfinance live (GC=F, Feb 2026)"],
            ["EUR/USD spot", "X\u2080", "1.181", "yfinance live (EURUSD=X, Feb 2026)"],
            ["USD risk-free rate", "r_USD", "3.6%", "13-week T-bill (^IRX)"],
            ["EUR risk-free rate", "r_EUR", "2.0%", "ECB deposit rate (configured)"],
            ["Gold volatility", "\u03c3_S", "41%", "EWMA (\u03bb=0.94, GC=F)"],
            ["EUR/USD volatility", "\u03c3_X", "6.2%", "EWMA (\u03bb=0.94, EURUSD=X)"],
            ["Correlation", "\u03c1", "\u22120.30", "126-day rolling (fallback)"],
            ["Gold convenience yield", "q", "3.6%", "Futures term structure (GC=F vs GCJ26.CMX)"],
        ],
        col_widths=[1.8, 0.8, 1.2, 2.7]
    )

    # 2.3 Risk-Neutral Valuation
    add_heading_numbered(doc, "2.3", "Risk-Neutral Valuation", level=2)

    add_body_text(doc, "The present value under the EUR risk-neutral measure:")
    add_formula(doc, "V\u2080 = \u1d3c\u1d60[e^(\u2212r_EUR \u00b7 \u03c4) \u00b7 Payoff\u03c4]")

    add_body_text(doc,
        "The path-dependent barrier feature precludes closed-form solutions, "
        "necessitating Monte Carlo methods.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 3: NUMERICAL IMPLEMENTATION
    # =========================================================================
    add_heading_numbered(doc, "3", "Numerical Implementation")

    add_heading_numbered(doc, "3.1", "Simulation Methodology", level=2)

    add_body_text(doc,
        "We employ Monte Carlo simulation with the following specifications:")

    add_styled_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["Simulation paths", "100,000", "Adequate precision for indicative pricing"],
            ["Time steps", "504", "Daily monitoring over 2 years"],
            ["Random seed", "Fixed", "Reproducibility"],
            ["Variance reduction", "Antithetic + Control", "~60% error reduction"],
        ],
        col_widths=[1.8, 1.5, 3.2]
    )

    doc.add_paragraph()
    add_body_text(doc, "Discretization Scheme", bold=True)
    add_body_text(doc, "Asset prices are simulated using the exact log-normal solution:")
    add_formula(doc,
        "S\u209c\u208a\u0394\u209c = S\u209c \u00b7 exp[(\u03bc_S \u2212 \u03c3_S\u00b2/2)\u0394t + \u03c3_S\u221a\u0394t \u00b7 Z\u2081]")
    add_formula(doc,
        "X\u209c\u208a\u0394\u209c = X\u209c \u00b7 exp[(\u03bc_X \u2212 \u03c3_X\u00b2/2)\u0394t + \u03c3_X\u221a\u0394t \u00b7 Z\u2082]")

    add_body_text(doc,
        "where Z\u2081, Z\u2082 are correlated standard normals generated via Cholesky decomposition:")
    add_formula(doc, "Z\u2082 = \u03c1 Z\u2081 + \u221a(1\u2212\u03c1\u00b2) Z\u22a5")

    add_heading_numbered(doc, "3.2", "Variance Reduction", level=2)

    add_body_text(doc, "Two techniques are implemented to improve computational efficiency:")

    add_body_text(doc,
        "Antithetic Variates: For each path with innovations {Z\u209c}, we also simulate the "
        "reflected path {\u2212Z\u209c}. The negative correlation between paired paths reduces variance.",
        bold=False)

    add_body_text(doc,
        "Control Variate: The vanilla gold forward (without barriers) serves as a control:")
    add_formula(doc,
        "V\u0302_adj = V\u0302_exotic + \u03b2(V_vanilla^analytical \u2212 V\u0302_vanilla)")
    add_body_text(doc,
        "where \u03b2 is the optimal control coefficient estimated from sample covariance.")

    add_body_text(doc,
        "Combined, these techniques reduce standard errors by approximately 60%.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 4: PRICING RESULTS
    # =========================================================================
    add_heading_numbered(doc, "4", "Pricing Results")

    # 4.1 Base Case
    add_heading_numbered(doc, "4.1", "Base Case Valuation", level=2)

    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Z Group Present Value", "EUR +63,968,230"],
            ["Alphabank Present Value", "EUR \u221263,968,230"],
            ["Standard Error", "EUR 896,173"],
            ["95% Confidence Interval", "[+62.2M, +65.7M]"],
        ],
        col_widths=[3.0, 3.5]
    )

    # 4.2 Barrier Analysis
    doc.add_paragraph()
    add_heading_numbered(doc, "4.2", "Barrier Analysis", level=2)

    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Overall Knock-Out Rate", "64.2%"],
            ["Lower Barrier Breaches", "27.3%"],
            ["Upper Barrier Breaches", "36.9%"],
            ["Average Time to Knock-Out", "0.98 years (11.8 months)"],
        ],
        col_widths=[3.0, 3.5]
    )

    doc.add_paragraph()
    add_body_text(doc,
        "The barrier breach profile has shifted dramatically compared to earlier market conditions. "
        "The upper barrier now accounts for the majority of knock-outs (37% vs 27% lower), reflecting "
        "the EUR/USD spot at 1.181\u2014only 5.8% from the upper barrier versus 11.1% from the lower. "
        "The interest rate differential (r_EUR \u2212 r_USD = \u22121.6%) still implies euro depreciation "
        "drift, but the proximity asymmetry dominates.")

    add_chart(doc, "gold_paths.png")
    add_chart(doc, "eurusd_paths.png")

    # 4.3 Convergence
    add_heading_numbered(doc, "4.3", "Convergence Verification", level=2)

    add_body_text(doc, "Monte Carlo estimates stabilize as path counts increase:")

    add_styled_table(doc,
        ["Paths", "Price Estimate", "Standard Error"],
        [
            ["5,000", "EUR +61.5M", "EUR 3,844K"],
            ["10,000", "EUR +61.6M", "EUR 2,698K"],
            ["25,000", "EUR +62.7M", "EUR 1,768K"],
            ["50,000", "EUR +63.1M", "EUR 1,241K"],
            ["100,000", "EUR +64.0M", "EUR 896K"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    doc.add_paragraph()
    add_body_text(doc,
        "Standard errors decay proportionally to 1/\u221an, confirming proper convergence behavior.")

    add_chart(doc, "convergence.png")

    doc.add_page_break()

    # =========================================================================
    # SECTION 5: CRITICAL ASSESSMENT
    # =========================================================================
    add_heading_numbered(doc, "5", "Critical Assessment")

    # 5.1 Strike Analysis
    add_heading_numbered(doc, "5.1", "Strike Price Analysis", level=2)

    add_body_text(doc,
        "The specified strike of USD 4,600/oz warrants careful examination.")

    add_body_text(doc, "Forward Price Calculation:", bold=True)
    add_formula(doc,
        "F\u2080,T = S\u2080 \u00b7 e^((r_USD \u2212 q) \u00b7 T) = 5,203 \u00b7 e^((0.036\u22120.036)\u00b72) \u2248 USD 5,203/oz")

    add_body_text(doc,
        "With convenience yield approximately equal to the USD risk-free rate, the forward is near "
        "spot. The strike sits 12% below the forward, placing Z Group in the money:")
    add_formula(doc,
        "Moneyness = F\u2080,T / K = 5,203 / 4,600 = 113.1%")

    add_body_text(doc, "Alternative Strike Analysis:", bold=True)
    add_styled_table(doc,
        ["Strike", "Forward Relationship", "Z Group PV"],
        [
            ["USD 4,000", "23% below forward", "EUR +146M"],
            ["USD 4,300", "17% below forward", "EUR +101M"],
            ["USD 4,600", "12% below forward", "EUR +63M"],
            ["USD 4,900", "6% below forward", "EUR +29M"],
            ["USD 5,200", "At-the-money", "EUR ~0M"],
            ["USD 5,500", "6% above forward", "EUR \u221227M"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    add_chart(doc, "gold_sensitivity.png")

    # 5.2 Barrier Configuration
    add_heading_numbered(doc, "5.2", "Barrier Configuration", level=2)

    add_body_text(doc,
        "With EUR/USD at 1.181, the upper barrier at 1.25 is now the proximate risk:")

    add_formula(doc,
        "Distance to Upper Barrier = (U \u2212 X\u2080) / X\u2080 = (1.25 \u2212 1.181) / 1.181 = 5.8%")
    add_formula(doc,
        "Distance to Lower Barrier = (X\u2080 \u2212 L) / X\u2080 = (1.181 \u2212 1.05) / 1.181 = 11.1%")

    add_body_text(doc,
        "With 6.2% annual EUR/USD volatility, the asymmetric positioning creates an "
        "upper-barrier-dominated knock-out profile.")

    add_body_text(doc, "Alternative Configurations:", bold=True)
    add_styled_table(doc,
        ["Corridor", "Knock-Out Rate", "Expected Duration"],
        [
            ["[1.05, 1.25]", "64%", "12 months"],
            ["[1.00, 1.30]", "27%", "15 months"],
            ["[0.95, 1.35]", "10%", "17 months"],
        ],
        col_widths=[2.0, 2.0, 2.5]
    )

    add_chart(doc, "fx_sensitivity.png")
    add_chart(doc, "scenario_analysis.png")

    doc.add_page_break()

    # =========================================================================
    # SECTION 6: RISK SENSITIVITIES
    # =========================================================================
    add_heading_numbered(doc, "6", "Risk Sensitivities")

    # 6.1 Greeks
    add_heading_numbered(doc, "6.1", "Greeks Summary", level=2)

    add_styled_table(doc,
        ["Greek", "Value", "Interpretation"],
        [
            ["Delta_gold", "EUR 105,677 per USD 1", "First-order gold sensitivity"],
            ["Gamma_gold", "EUR \u2212669", "Gold convexity"],
            ["Delta_FX", "EUR \u22128.18M per 0.01 FX", "EUR/USD sensitivity (negative: upper barrier risk)"],
            ["Vega_gold", "EUR \u22122.88M per 1% vol", "Gold vega (higher vol increases KO probability)"],
            ["Rho_EUR", "EUR \u22121,031M per 1bp", "EUR rate sensitivity"],
            ["Rho_corr", "EUR \u2212973K per 0.05 corr", "Correlation sensitivity"],
        ],
        col_widths=[1.2, 2.5, 2.8]
    )

    add_chart(doc, "greeks_chart.png")

    # 6.2 Hedging
    add_heading_numbered(doc, "6.2", "Hedging Implications", level=2)

    add_body_text(doc,
        "Delta Hedging: The gold delta of EUR 106K per dollar implies a hedge ratio of approximately:",
        bold=False)
    add_formula(doc,
        "Gold Hedge = (\u0394_gold / S\u2080) \u00d7 K = (105,677 / 5,203) \u00d7 4,600 \u2248 93,400 oz")

    add_body_text(doc,
        "Barrier Risk: With EUR/USD at 1.181, the upper barrier at 1.25 is only 5.8% away. "
        "As EUR/USD approaches either barrier, gamma and delta become increasingly unstable\u2014"
        "the characteristic \u201cpin risk\u201d of barrier options. The negative FX delta "
        "(EUR \u22128.2M per 0.01) reflects that EUR appreciation toward 1.25 destroys contract "
        "value through knock-out.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 7: MODEL VALIDATION
    # =========================================================================
    add_heading_numbered(doc, "7", "Model Validation")

    # 7.1 Model Comparison
    add_heading_numbered(doc, "7.1", "Alternative Specifications", level=2)

    add_body_text(doc,
        "To ensure robustness, we compared valuations across three model specifications:")

    add_styled_table(doc,
        ["Model", "Z Group PV", "Knock-Out Rate"],
        [
            ["Base GBM", "EUR +63.1M", "64.5%"],
            ["Heston Stochastic Vol", "EUR +64.1M", "64.1%"],
            ["Merton Jump-Diffusion", "EUR +65.7M", "64.1%"],
        ],
        col_widths=[2.5, 2.0, 2.0]
    )

    doc.add_paragraph()
    add_body_text(doc,
        "All models converge within 4%, with Heston and Merton producing slightly higher "
        "valuations due to stochastic volatility and jump dynamics amplifying the in-the-money "
        "payoff. Model specification risk remains secondary to market parameter uncertainty.")

    # 7.2 Analytical Benchmark
    add_heading_numbered(doc, "7.2", "Analytical Benchmark", level=2)

    add_body_text(doc,
        "The vanilla gold forward (without barriers) provides a sanity check:")
    add_formula(doc,
        "V_vanilla = e^(\u2212r_EUR \u00b7 T) \u00b7 N \u00b7 (F\u2080,T \u2212 K) / K "
        "= e^(\u22120.020 \u00d7 2) \u00b7 500M \u00d7 (5,203 \u2212 4,600) / 4,600")
    add_formula(doc,
        "V_vanilla = EUR +63,004,791")

    add_body_text(doc,
        "The knock-out version (EUR +64.0M) is EUR 1.0M higher than the vanilla, reflecting "
        "that early termination locks in profits when gold is above strike\u2014a reversal from "
        "the prior regime where knock-outs destroyed value.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 8: CONCLUSIONS
    # =========================================================================
    add_heading_numbered(doc, "8", "Conclusions and Recommendations")

    add_heading_numbered(doc, "8.1", "Summary of Findings", level=2)

    add_body_text(doc,
        "The proposed structure is technically sound and priceable using standard Monte Carlo "
        "techniques. The February 2026 market environment has fundamentally altered the risk profile:")

    add_body_text(doc,
        "1.  Position reversal: Gold\u2019s surge to $5,203/oz places Z Group firmly in the money "
        "(113% moneyness), with a positive PV of EUR +64M. This contrasts sharply with earlier "
        "market conditions where the position was deeply out-of-the-money.")

    add_body_text(doc,
        "2.  Barrier risk shift: The upper EUR/USD barrier at 1.25 is now only 5.8% from spot, "
        "making it the primary knock-out driver (37% of paths). The 64% total knock-out rate and "
        "12-month expected duration represent a more balanced risk profile than the prior 93%/5-month "
        "scenario.")

    add_body_text(doc,
        "3.  Elevated volatility: Gold volatility at 41% (EWMA) is more than double historical "
        "norms, increasing both the potential upside for Z Group and the mark-to-market volatility "
        "for Alphabank\u2019s hedging book.")

    add_heading_numbered(doc, "8.2", "Recommendations", level=2)

    add_body_text(doc,
        "We recommend proceeding to term sheet stage with the following considerations:")

    add_body_text(doc,
        "1.  Alphabank should carefully assess credit exposure given the positive Z Group PV, "
        "ensuring adequate collateral arrangements")

    add_body_text(doc,
        "2.  The upper barrier proximity (5.8%) warrants active monitoring\u2014EUR appreciation "
        "events could trigger knock-out and crystallize Z Group\u2019s gain")

    add_body_text(doc,
        "3.  Consider whether alternative barrier configurations (e.g., [1.00, 1.30] with 27% "
        "KO rate) would provide Z Group with more durable hedging protection")

    add_heading_numbered(doc, "8.3", "Next Steps", level=2)

    add_body_text(doc, "Upon Committee approval, we will:")
    for item in [
        "Finalize term sheet documentation",
        "Establish hedging framework with Trading Desk",
        "Coordinate credit approval with Risk Management",
        "Schedule client presentation",
    ]:
        add_body_text(doc, f"\u2022  {item}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("[END OF MEMORANDUM]")
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # =========================================================================
    # APPENDIX A: NOMENCLATURE
    # =========================================================================
    doc.add_heading("Appendix A: Nomenclature", level=1)

    add_styled_table(doc,
        ["Symbol", "Description"],
        [
            ["S\u209c", "Gold spot price (USD/oz) at time t"],
            ["X\u209c", "EUR/USD exchange rate at time t"],
            ["K", "Strike price (USD 4,600/oz)"],
            ["N", "Notional principal (EUR 500M)"],
            ["T", "Maturity (2 years)"],
            ["L, U", "Lower (1.05) and upper (1.25) barriers"],
            ["r_EUR, r_USD", "Risk-free rates"],
            ["\u03c3_S, \u03c3_X", "Volatilities"],
            ["\u03c1", "Correlation coefficient"],
            ["q", "Gold convenience yield"],
            ["\u03c4", "Settlement time"],
            ["\u03c4_KO", "Knock-out time"],
        ],
        col_widths=[1.5, 5.0]
    )

    doc.add_page_break()

    # =========================================================================
    # APPENDIX B: FIGURE INDEX
    # =========================================================================
    doc.add_heading("Appendix B: Figure Index", level=1)

    figures = [
        ("gold_paths.png", "Monte Carlo Simulation: Gold Price Paths"),
        ("eurusd_paths.png", "Monte Carlo Simulation: EUR/USD Paths with Barriers"),
        ("convergence.png", "Monte Carlo Convergence Analysis"),
        ("gold_sensitivity.png", "Gold Price Sensitivity Analysis"),
        ("fx_sensitivity.png", "EUR/USD Sensitivity Analysis"),
        ("scenario_analysis.png", "Scenario Analysis \u2014 Z Group PV Impact"),
        ("greeks_chart.png", "Risk Sensitivities (Greeks)"),
        ("volatility_surface.png", "Volatility Surface (2D Grid)"),
        ("correlation_sensitivity.png", "Correlation Sensitivity Analysis"),
        ("payoff_distribution.png", "Payoff Distribution"),
        ("knockout_analysis.png", "Knockout Time Distribution and Barrier Breach Analysis"),
        ("valuation_comparison.png", "Market Conditions Change (Jan \u2192 Feb 2026)"),
        ("gold_history.png", "Gold Price Evolution"),
    ]

    for i, (fname, title) in enumerate(figures, 1):
        add_body_text(doc, f"{i}.  {title} \u2014 {fname}")

    # =========================================================================
    # APPENDIX C: ADDITIONAL CHARTS
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("Appendix C: Additional Charts", level=1)

    add_body_text(doc,
        "The following charts provide supplementary analysis supporting the main body of this memorandum.")

    additional_charts = [
        ("volatility_surface.png", "Volatility Surface"),
        ("correlation_sensitivity.png", "Correlation Sensitivity"),
        ("payoff_distribution.png", "Payoff Distribution"),
        ("knockout_analysis.png", "Knockout Analysis"),
        ("valuation_comparison.png", "Valuation Comparison: Jan vs Feb 2026"),
        ("gold_history.png", "Gold Price History"),
    ]

    for fname, title in additional_charts:
        add_body_text(doc, title, bold=True)
        add_chart(doc, fname)
        doc.add_paragraph()

    return doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating PRODUCT_PROPOSAL_FINAL.docx ...")
    doc = build_document()

    output_path = "PRODUCT_PROPOSAL_FINAL.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")

    # Also copy to submission/
    os.makedirs("submission", exist_ok=True)
    import shutil
    shutil.copy2(output_path, "submission/PRODUCT_PROPOSAL_FINAL.docx")
    print(f"Copied to: submission/{output_path}")

    print("Done!")
